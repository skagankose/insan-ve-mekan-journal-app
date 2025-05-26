from docx import Document
from pathlib import Path
from typing import List, Optional, Dict
import os
import traceback # For detailed error logging
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH  # <-- Add this import
from docx.oxml import OxmlElement  # <-- Add this import
from docx.oxml.ns import qn        # <-- Add this import
from docx.shared import Inches, Pt  # Add this import for handling image dimensions and font size

def create_element(name: str) -> OxmlElement:
    """Create a new element with the given name."""
    return OxmlElement(name)

def create_attribute(element: OxmlElement, name: str, value: str) -> None:
    """Create an attribute on the given element."""
    element.set(qn(name), value)

def add_page_number(paragraph):
    """Add page number field to the given paragraph."""
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    run._r.append(fldChar1)

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar2)

def add_page_numbers(doc, skip_first_page: bool = False):
    """Add page numbers to the document."""
    section = doc.sections[0]
    footer = section.footer
    
    if skip_first_page:
        section.different_first_page_header_footer = True
    
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style.font.size = Pt(10)  # Set font size to 10pt
    add_page_number(paragraph)

def merge_docx_files(file_paths: List[str], output_path: str, cover_photo_path: Optional[str] = None) -> bool:
    """
    Merge multiple .docx files into a single document, optionally with a cover photo on the first page.
    Page numbers are added to all pages except the cover page.
    
    Args:
        file_paths: List of paths to .docx files to merge.
        output_path: Path where the merged .docx file will be saved.
        cover_photo_path: Optional path to a cover photo to be added as the first page.
        
    Returns:
        bool: True if successful, False otherwise.
    """
    print(f"Starting merge with cover photo and page numbers. Output path: {output_path}")
    print(f"Files to merge: {file_paths}")
    try:
        merged_doc = Document()
        # Remove the default paragraph often created in a new document if it's empty,
        # to ensure the first merged content starts cleanly.
        if len(merged_doc.paragraphs) == 1 and not merged_doc.paragraphs[0].text:
            p = merged_doc.paragraphs[0]
            p_element = p._element
            p_element.getparent().remove(p_element)

        # Add cover photo if provided
        has_cover = False
        if cover_photo_path and os.path.exists(cover_photo_path):
            print(f"Adding cover photo from: {cover_photo_path}")
            has_cover = True
            # Add the cover photo
            paragraph = merged_doc.add_paragraph()
            run = paragraph.add_run()
            # Add the image with specific dimensions
            run.add_picture(cover_photo_path, width=Inches(8.0))  # Adjust width as needed
            
            # Add a page break after the cover photo
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
        elif cover_photo_path:
            print(f"Cover photo not found at path: {cover_photo_path}")

        for file_idx, file_path in enumerate(file_paths):
            print(f"Processing file {file_idx + 1}/{len(file_paths)}: {file_path}")
            if not file_path or not os.path.exists(file_path):
                print(f"File not found or path is invalid: {file_path}. Skipping.")
                continue
                
            if not file_path.lower().endswith('.docx'):
                print(f"Skipping non-docx file: {file_path}")
                continue

            # Insert a page break as an XML element before appending content, except for the first file
            if file_idx > 0:
                page_break = OxmlElement('w:p')
                run = OxmlElement('w:r')
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run.append(br)
                page_break.append(run)
                merged_doc.element.body.append(page_break)

            source_doc = Document(file_path)
            print(f"Opened source document: {file_path}")
            
            appended_element_count = 0
            for element in source_doc.element.body:
                merged_doc.element.body.append(element)
                appended_element_count += 1
            print(f"Appended {appended_element_count} elements from {file_path}")

        # Add page numbers to all pages (skip first page if there's a cover photo)
        add_page_numbers(merged_doc, skip_first_page=has_cover)

        # Ensure the output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            print(f"Creating output directory: {output_dir}")
            os.makedirs(output_dir)
        
        print(f"Saving merged document to: {output_path}")
        merged_doc.save(output_path)
        print(f"Successfully merged files with cover photo and page numbers into {output_path}")
        return True
        
    except Exception as e:
        print(f"Error during merge_docx_files: {e}")
        traceback.print_exc() # Print full traceback for debugging
        return False 

def create_table_of_contents(entries: List[Dict], output_path: str) -> bool:
    """
    Create a table of contents document for journal entries.
    
    Args:
        entries: List of journal entries with title and other metadata
        output_path: Path where the table of contents file will be saved
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        doc = Document()
        
        # Add title
        title = doc.add_paragraph("Table of Contents")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Add a line break after title
        doc.add_paragraph()
        
        # Add entries
        for idx, entry in enumerate(entries, 1):
            # Create entry line with title
            p = doc.add_paragraph()
            p.add_run(f"{idx}. ").bold = True
            p.add_run(entry['title'])
            
            # Add Article Type and Page Number on the same line as title, right-aligned or spaced
            # For simplicity, adding it after the title. Proper right alignment is more complex.
            if entry.get('article_type') and entry.get('page_number'):
                p.add_run(f" ({entry.get('article_type', 'N/A')}, s. {entry.get('page_number', 'N/A')})").italic = True
            elif entry.get('article_type'):
                p.add_run(f" ({entry.get('article_type', 'N/A')})").italic = True
            elif entry.get('page_number'):
                 p.add_run(f" (s. {entry.get('page_number', 'N/A')})").italic = True

            # Add authors if available
            if entry.get('authors'):
                try:
                    author_para = doc.add_paragraph()
                    author_para.paragraph_format.left_indent = Inches(0.5)
                    author_para.add_run("Authors: ").italic = True
                    # Ensure authors is treated as string
                    authors_text = str(entry['authors']) if entry['authors'] else ''
                    author_para.add_run(authors_text)
                except Exception as e:
                    print(f"Warning: Could not add authors for entry {idx}: {e}")
                    # Continue to next part if authors can't be added, do not skip the whole entry
            
            # Add abstract if available (first few words)
            abstract = entry.get('abstract_tr') or entry.get('abstract_en')
            if abstract:
                abstract_preview = abstract[:200] + "..." if len(abstract) > 200 else abstract
                abstract_para = doc.add_paragraph()
                abstract_para.paragraph_format.left_indent = Inches(0.5)
                abstract_para.add_run("Abstract: ").italic = True
                abstract_para.add_run(abstract_preview)
            
            # Add a line break between entries
            doc.add_paragraph()
        
        # Add page numbers
        add_page_numbers(doc)
        
        # Create output directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save the document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating table of contents: {e}")
        traceback.print_exc()
        return False 