from docx import Document
from pathlib import Path
from typing import List, Optional, Dict
import os
import traceback # For detailed error logging
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH  # <-- Add this import
from docx.oxml import OxmlElement  # <-- Add this import
from docx.oxml.ns import qn        # <-- Add this import
from docx.shared import Inches, Pt  # Add this import for handling image dimensions and font size
from docxcompose.composer import Composer # Add this import

def create_element(name: str) -> OxmlElement:
    """Create a new element with the given name."""
    return OxmlElement(name)

def create_attribute(element: OxmlElement, name: str, value: str) -> None:
    """Create an attribute on the given element."""
    element.set(qn(name), value)

def add_page_number(paragraph):
    """Add page number field to the given paragraph."""
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    run._r.append(fldChar1)

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar2)

def add_page_numbers(doc, skip_first_page: bool = False):
    """Add page numbers to the document."""
    section = doc.sections[0]
    footer = section.footer
    
    if skip_first_page:
        section.different_first_page_header_footer = True
    
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style.font.size = Pt(10)  # Set font size to 10pt
    add_page_number(paragraph)

def merge_docx_files(file_paths: List[str], output_path: str, cover_photo_path: Optional[str] = None) -> bool:
    """
    Merge multiple .docx files into a single document, optionally with a cover photo on the first page.
    Page numbers are added to all pages except the cover page.
    
    Args:
        file_paths: List of paths to .docx files to merge.
        output_path: Path where the merged .docx file will be saved.
        cover_photo_path: Optional path to a cover photo to be added as the first page.
        
    Returns:
        bool: True if successful, False otherwise.
    """
    print(f"Starting merge with cover photo and page numbers. Output path: {output_path}")
    print(f"Files to merge: {file_paths}")
    try:
        # Initialize the first document, which will be the base for the Composer
        # If a cover photo is provided, it will be handled separately first,
        # and then the rest of the documents will be appended.
        
        master_doc: Optional[Document] = None
        composer: Optional[Composer] = None
        initial_files_to_merge = list(file_paths) # Create a mutable copy

        # Ensure the output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            print(f"Creating output directory: {output_dir}")
            os.makedirs(output_dir)

        has_cover = False
        if cover_photo_path and os.path.exists(cover_photo_path):
            print(f"Adding cover photo from: {cover_photo_path}")
            has_cover = True
            # Create a new document for the cover photo
            cover_doc = Document()
            paragraph = cover_doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(cover_photo_path, width=Inches(8.0)) # Adjust width as needed
            run.add_break(WD_BREAK.PAGE)
            
            # Save the cover doc temporarily to be used by Composer
            temp_cover_path = os.path.join(output_dir, "temp_cover_for_merge.docx")
            cover_doc.save(temp_cover_path)
            
            master_doc = Document(temp_cover_path)
            composer = Composer(master_doc)

            # If there are other files to merge, append them after the cover
            if initial_files_to_merge:
                 # The first "real" document doesn't need a preceding page break
                first_real_doc_path = initial_files_to_merge.pop(0)
                if os.path.exists(first_real_doc_path) and first_real_doc_path.lower().endswith('.docx'):
                    print(f"Appending first document after cover: {first_real_doc_path}")
                    source_doc = Document(first_real_doc_path)
                    composer.append(source_doc)
                else:
                    print(f"Skipping invalid or non-docx file: {first_real_doc_path}")
            
            # Clean up temporary cover file
            # os.remove(temp_cover_path) # Let's keep it for now for debugging, remove later if all good.

        elif initial_files_to_merge:
            # No cover photo, start with the first document in the list
            first_doc_path = initial_files_to_merge.pop(0)
            if os.path.exists(first_doc_path) and first_doc_path.lower().endswith('.docx'):
                print(f"Setting first document as master: {first_doc_path}")
                master_doc = Document(first_doc_path)
                composer = Composer(master_doc)
            else:
                print(f"Skipping invalid or non-docx file as first document: {first_doc_path}")
                # If the very first document is invalid, we can't initialize Composer
                # Or, we could try to find the *next* valid one. For now, let's assume this is an error or edge case.
                # If all files are invalid, this will fail later.
                # For robustness, we should initialize composer with an empty doc if first_doc_path is bad AND there are other files
                if not initial_files_to_merge : # No other files left
                     print(f"Error: First document is invalid and no other documents to merge: {first_doc_path}")
                     return False # Or handle by creating an empty merged doc

        if not composer:
            # This happens if there was no cover and the first file in initial_files_to_merge was invalid
            # OR if initial_files_to_merge was empty from the start (and no cover)
            if not initial_files_to_merge and not has_cover : # no files were provided at all
                print("No files provided to merge and no cover photo. Creating an empty document.")
                merged_doc = Document() # Create an empty document
                # Remove default paragraph if it's empty
                if len(merged_doc.paragraphs) == 1 and not merged_doc.paragraphs[0].text:
                    p = merged_doc.paragraphs[0]
                    p_element = p._element
                    p_element.getparent().remove(p_element)
                # Add page numbers if needed (though for an empty doc, it's likely not)
                # add_page_numbers(merged_doc, skip_first_page=False)
                merged_doc.save(output_path)
                print(f"Saved an empty document to: {output_path}")
                return True # Successfully "merged" an empty list of files
            elif not initial_files_to_merge and has_cover and master_doc: # Only cover was provided
                print("Only a cover photo was provided. Saving document with only cover.")
                # Page numbers for a single cover page might not be desired.
                # If add_page_numbers is called, it should respect skip_first_page=True
                add_page_numbers(master_doc, skip_first_page=True) # master_doc here is the cover_doc
                master_doc.save(output_path)
                # if os.path.exists(temp_cover_path): os.remove(temp_cover_path) # cleanup
                print(f"Successfully saved document with only cover photo to {output_path}")
                return True

            # If we reach here, it means the first file was bad, but there might be others.
            # Initialize with an empty doc and try to append the rest.
            print("First document was invalid or no initial files. Initializing with an empty document for subsequent appends.")
            master_doc = Document()
            # Remove the default paragraph if it's empty
            if len(master_doc.paragraphs) == 1 and not master_doc.paragraphs[0].text:
                p = master_doc.paragraphs[0]
                p_element = p._element
                p_element.getparent().remove(p_element)
            composer = Composer(master_doc)
            # has_cover is False here, as this block is in the elif of `if cover_photo_path...`
            # or the first doc in initial_files_to_merge was bad.

        # Append the rest of the documents
        for file_idx, file_path in enumerate(initial_files_to_merge):
            print(f"Processing file {file_idx + (1 if has_cover or master_doc.element.body else 0) +1}/{len(file_paths)}: {file_path}")
            if not file_path or not os.path.exists(file_path):
                print(f"File not found or path is invalid: {file_path}. Skipping.")
                continue
            
            if not file_path.lower().endswith('.docx'):
                print(f"Skipping non-docx file: {file_path}")
                continue

            # Add a page break before appending the new document's content
            # Composer handles document merging, but explicit page breaks between appended docs are good.
            # The first document (or cover) doesn't need a preceding break.
            # Subsequent documents appended via composer.append() will start on a new page if the previous doc ended properly.
            # However, to be absolutely sure, we can add a page break to the *composer's current document*
            # *before* appending the next one. This is tricky with composer.
            # A simpler approach for docxcompose is that it usually handles the separation.
            # If not, one might need to add a page break to the *end* of the previous `source_doc`
            # before `composer.append(source_doc)`, or add a paragraph with a page break to the `master_doc`
            # *before* calling `composer.append(source_doc)`.

            # Let's try without explicit page breaks first, as composer might handle it.
            # If sections run together, we'll add manual page breaks.
            # Typically, composer.append starts the new content on a new page by default if the
            # previous content ended with a section break or if a new section is started.
            # For simple appends, it might just flow.
            # Let's add an explicit page break to the master document *before* appending.
            if composer.doc.element.body: # Ensure there's content to add a break after
                 # Check if the last element is already a page break, to avoid double breaks
                last_element = composer.doc.element.body[-1]
                is_last_element_page_break = False
                if last_element.tag == qn('w:p'):
                    for run_element in last_element.findall(qn('w:r')):
                        if run_element.find(qn('w:br')) is not None and run_element.find(qn('w:br')).get(qn('w:type')) == 'page':
                            is_last_element_page_break = True
                            break
                if not is_last_element_page_break:
                    composer.doc.add_page_break()
            
            print(f"Appending document: {file_path}")
            source_doc = Document(file_path)
            composer.append(source_doc)
            print(f"Successfully appended {file_path}")

        # Get the final merged document from the composer
        # The composer modifies master_doc in place.
        merged_doc_final = composer.doc 

        # Add page numbers to all pages (skip first page if there's a cover photo)
        add_page_numbers(merged_doc_final, skip_first_page=has_cover)
        
        print(f"Saving merged document to: {output_path}")
        composer.save(output_path)  # This preserves headers, footers, and sections!

        # Clean up temporary cover file if it was created
        if has_cover and 'temp_cover_path' in locals() and os.path.exists(temp_cover_path):
            try:
                os.remove(temp_cover_path)
                print(f"Removed temporary cover file: {temp_cover_path}")
            except Exception as e:
                print(f"Warning: Could not remove temporary cover file {temp_cover_path}: {e}")

        print(f"Successfully merged files with cover photo and page numbers into {output_path}")
        return True
    except Exception as e:
        print(f"Error during merge_docx_files: {e}")
        traceback.print_exc() # Print full traceback for debugging
        return False

def create_table_of_contents(entries: List[Dict], output_path: str) -> bool:
    """
    Create a table of contents document for journal entries.
    
    Args:
        entries: List of journal entries with title and other metadata
        output_path: Path where the table of contents file will be saved
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        doc = Document()
        
        # Add title
        title = doc.add_paragraph("İçindekiler")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Add a line break after title
        doc.add_paragraph()
        
        # Add entries
        for idx, entry in enumerate(entries, 1):
            # Create entry line with title
            p = doc.add_paragraph()
            p.add_run(f"{idx}. ").bold = True
            p.add_run(entry['title'])
            
            # Add Article Type and Page Number on the same line as title, right-aligned or spaced
            # For simplicity, adding it after the title. Proper right alignment is more complex.
            '''
            if entry.get('article_type') and entry.get('page_number'):
                p.add_run(f" ({entry.get('article_type', 'N/A')}, s. {entry.get('page_number', 'N/A')})").italic = True
            elif entry.get('article_type'):
                p.add_run(f" ({entry.get('article_type', 'N/A')})").italic = True
            elif entry.get('page_number'):
                 p.add_run(f" (s. {entry.get('page_number', 'N/A')})").italic = True
            '''
            if entry.get('page_number'):
                p.add_run(f" (s. {entry.get('page_number', 'DSB')})").italic = True

            # Add authors if available
            if entry.get('authors'):
                try:
                    author_para = doc.add_paragraph()
                    author_para.paragraph_format.left_indent = Inches(0.5)
                    author_para.add_run("Yazarlar: ").italic = True
                    # Ensure authors is treated as string
                    authors_text = str(entry['authors']) if entry['authors'] else ''
                    author_para.add_run(authors_text)
                except Exception as e:
                    print(f"Warning: Could not add authors for entry {idx}: {e}")
                    # Continue to next part if authors can't be added, do not skip the whole entry
            
            # Add abstract if available (first few words)
            abstract = entry.get('abstract_tr') or entry.get('abstract_en')
            if abstract:
                abstract_preview = abstract[:200] + "..." if len(abstract) > 200 else abstract
                abstract_para = doc.add_paragraph()
                abstract_para.paragraph_format.left_indent = Inches(0.5)
                abstract_para.add_run("Özet: ").italic = True
                abstract_para.add_run(abstract_preview)
            
            # Add a line break between entries
            doc.add_paragraph()
        
        # Add page numbers
        add_page_numbers(doc)
        
        # Create output directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save the document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating table of contents: {e}")
        traceback.print_exc()
        return False 